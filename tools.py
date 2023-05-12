#!/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File    :    Untitled-1
@Time    :    2023/05/04 16:25:41
@Author  :    cyq
@Version :    1.0
@Contact :    1135362921@qq.com
@Desc    :    
'''

import sys
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import openpyxl as xl
import xlwings as xw
import random
import math
from win32com import client as wc
import win32print
import tempfile
import win32api
import pythoncom
import time
from natsort import natsorted


def log(*msg):
    f=open('日志.txt','a+',encoding='UTF-8')
    f.write(f'{time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())}  :  {msg}\n')
    f.close()
    
class Tool:

    def zhuanhuan(path,dest_path,*args):
        print(args)
        if args==():
            args=('docx',16,'xlsx',51)
        pythoncom.CoInitialize()
        word= wc.Dispatch("Word.application")
        word.Visible =0
        word.DisplayAlerts =0
        excel=wc.gencache.EnsureDispatch('Excel.Application')
        for file in os.listdir(path):
            print('file:',file)
           
            if os.path.isfile(path +'\\'+ file) and os.path.splitext(file)[1] in ['.doc','.docx']:
                log("文件名",file)
                (file_path, temp_file_name) = os.path.split(file)
                (short_name, extension) = os.path.splitext(temp_file_name)
                doc = word.Documents.Open(path +'\\'+ file)
                doc.SaveAs(dest_path +'\\'+ short_name + f".{args[0]}", args[1])
                doc.Close()
            elif os.path.isfile(path +'\\'+ file) and os.path.splitext(file)[1] in ['.xls','.xlsx']:
                log("文件名",file)
                (file_path, temp_file_name) = os.path.split(file)
                (short_name, extension) = os.path.splitext(temp_file_name)
                wb = excel.Workbooks.Open(path +'\\'+ file)
                if args[2]!='pdf':
                    wb.SaveAs(dest_path +'\\'+ short_name + f".{args[2]}", args[3])
                else:
                    wb.ExportAsFixedFormat(args[3], dest_path +'\\'+ short_name + f".{args[2]}")
                wb.Close()  
            
        word.Quit()
        excel.Quit()
        pythoncom.CoUninitialize()

    def replace_word(old_info, new_info,document):
        
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if run.text and old_info in run.text:
                    rt=run.text.replace(old_info, new_info)
                    run.text=rt

        for table in document.tables:
            for row in table.rows:
                 for cell in row.cells:
                     for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text and old_info in run.text:
                                rt=run.text.replace(old_info, new_info)
                                run.text=rt
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    if run.text and old_info in run.text:
                        rt=run.text.replace(old_info, new_info)
                        run.text=rt
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    if run.text and old_info in run.text:
                        rt=run.text.replace(old_info, new_info)
                        run.text=rt                        
    #xlwings
    def replace_excel(old,new,wb):    
        for sht in wb.sheets:
            cell = sht.used_range.last_cell
            rows = cell.row
            columns = cell.column
            dt=sht.range((1,1),(rows,columns)).value
            if dt==None:
                return None
            for row in dt:
                for cv in row:
                    if cv !=None:
                        if isinstance(cv,str) and old in cv:
                            i=dt.index(row)+1
                            j=row.index(cv)+1
                            sht.range(i,j).value=cv.replace(old,new)

                    
                    
    '''
    打印文件
    '''
    def Printer(filename):
        open(filename, "r")
        win32api.ShellExecute(
            0,
            "print",
            filename,
            '/d:"%s"' % win32print.GetDefaultPrinter(),
            ".",
            0
        )